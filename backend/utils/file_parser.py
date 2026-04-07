"""
File Parser — extracts and cleans text from PDF, DOCX, and TXT files.

Pipeline:
  1. Parse raw bytes → raw text
  2. Normalize: fix encoding artifacts, collapse whitespace, strip noise
  3. Validate: reject empty or too-short output
"""
import io
import re


# ── Public API ────────────────────────────────────────────────────────────────

def parse_file(content: bytes, ext: str) -> str:
    """
    Parse file bytes → cleaned plain text.
    Raises ValueError on bad input, ImportError if a library is missing.
    """
    ext = ext.lower().strip()
    parsers = {
        ".txt":  _parse_txt,
        ".pdf":  _parse_pdf,
        ".docx": _parse_docx,
        ".doc":  _parse_docx,
    }
    if ext not in parsers:
        raise ValueError(f"Unsupported file type '{ext}'. Accepted: PDF, TXT, DOCX.")

    raw = parsers[ext](content)
    cleaned = _clean_text(raw)

    if not cleaned or len(cleaned) < 30:
        raise ValueError(
            "Document contains no usable text after parsing. "
            "If it is a scanned PDF, please use a text-based version."
        )
    return cleaned


# ── Parsers ───────────────────────────────────────────────────────────────────

def _parse_txt(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1", "cp1252"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode text file — unknown encoding.")


def _parse_pdf(content: bytes) -> str:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ImportError("PyPDF2 is not installed. Run: pip install PyPDF2")

    reader = PdfReader(io.BytesIO(content))
    if not reader.pages:
        raise ValueError("PDF has no pages.")

    pages = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text)
        except Exception:
            continue  # skip unreadable pages rather than crashing

    if not pages:
        raise ValueError(
            "PDF contains no extractable text. "
            "It may be a scanned image — please use a text-based PDF."
        )
    return "\n\n".join(pages)


def _parse_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    doc = Document(io.BytesIO(content))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    if not parts:
        raise ValueError("DOCX file contains no readable text content.")
    return "\n\n".join(parts)


# ── Text Cleaning ─────────────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """
    Normalize extracted text for clean AI input.
    Steps applied in order:
      1. Fix common unicode/mojibake artifacts
      2. Remove null bytes and non-printable control characters
      3. Normalize line endings
      4. Fix hyphenated line-breaks (PDF word-wrap artifact)
      5. Collapse excessive blank lines (max 2 consecutive)
      6. Collapse internal whitespace within lines
      7. Remove standalone page-number lines
    """
    if not text:
        return ""

    # 1. Fix common PDF mojibake
    text = _fix_mojibake(text)

    # 2. Remove null bytes and non-printable control chars (keep \n, \t)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

    # 3. Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # 4. Fix PDF hyphenated line-breaks: "hyphen-\nated" → "hyphenated"
    text = re.sub(r"-\n([a-z])", r"\1", text)

    # 5 & 6. Process line by line
    lines = [line.rstrip() for line in text.split("\n")]
    cleaned_lines = []
    blank_count = 0
    for line in lines:
        if line == "":
            blank_count += 1
            if blank_count <= 2:
                cleaned_lines.append(line)
        else:
            blank_count = 0
            cleaned_lines.append(re.sub(r"[ \t]{2,}", " ", line))

    text = "\n".join(cleaned_lines).strip()

    # 7. Remove page number artifacts
    text = _remove_page_artifacts(text)

    return text.strip()


def _fix_mojibake(text: str) -> str:
    """Fix common PDF/DOCX encoding artifacts."""
    replacements = {
        "\u2019": "'",   # right single quotation mark
        "\u2018": "'",   # left single quotation mark
        "\u201c": '"',   # left double quotation mark
        "\u201d": '"',   # right double quotation mark
        "\u2013": "-",   # en dash
        "\u2014": "--",  # em dash
        "\u2026": "...", # ellipsis
        "\u00a0": " ",   # non-breaking space
        "\ufb01": "fi",  # fi ligature
        "\ufb02": "fl",  # fl ligature
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    return text


def _remove_page_artifacts(text: str) -> str:
    """Remove standalone page numbers and other PDF/DOCX layout noise."""
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        stripped = line.strip()
        # Skip lines that are just a page number like "1", "Page 2", "- 3 -"
        if re.fullmatch(r"[-\s]*(page\s*)?\d{1,4}[-\s]*", stripped, re.IGNORECASE):
            continue
        cleaned.append(line)
    return "\n".join(cleaned)
